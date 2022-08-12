import docx
import os
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from extract_Alcance_PH import *
from numero_letras import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
import itertools

dictEnum = {1:"Primera (I)",2:"Segunda (II)",3:"Tercera (III)",4:"Cuarta (IV)",5:"Quinta (V)"}








dict_letters = {}

listofdatt , tipomod , dicdeff, dicTOT,manzana= getVals(1)

codUb = listofdatt[0][4]
ubGeneral = listofdatt[3][1]
descrGeneral = listofdatt[6][0]

faSe = listofdatt[0][5]
uniTotal = 0
for i in dicTOT:
    uniTotal += dicTOT[i]
def getTot():
    return uniTotal
ubDeCod = listofdatt[0][1]
valTerr = listofdatt[0][2]
NdeFinca = listofdatt[0][3]
pH = listofdatt[0][0]
aTotTerr = str(MakeSQRmtrr(listofdatt[0][7]))+" mts2"
aTotConst = str(MakeSQRmtrr(listofdatt[0][8]))+" mts2"
aTotCom = str(MakeSQRmtrr(listofdatt[0][9]))+" mts2"
aResComer = str(MakeSQRmtrr(listofdatt[3][0]))+" mts2"
aRes = str(MakeSQRmtrr(listofdatt[1][10]))+" mts2"
totFase = str(MakeSQRmtrr(listofdatt[1][7]))+" mts2"
constFase = str(MakeSQRmtrr(listofdatt[1][8]))+" mts2"
comFase = str(MakeSQRmtrr(listofdatt[1][9]))+" mts2"

for p in dicdeff:


    x = p.split()
    if len(x)> 2:
        for i in range(len(x)-2):
            i +=2
            x[1] =x[1]+ " "+x[i]
    if x[0] not in dict_letters:
        dict_letters[x[0]] = []
    
    if len(x) != 1:
        dict_letters[x[0]].append(x[1])


def getDDict(tipomod):
    list_of_phase = []
    for i in tipomod:
        for p in tipomod[i]:
            list_of_phase.append(p)
    totcount_phase = len(list_of_phase)
    

    dictOfLists = {}
    
    for p in list_of_phase:
        ind = p.find('-')
        if p[:ind] not in dictOfLists:
            dictOfLists[p[0:ind]] = [int(p[ind+1:])]
        else:
            dictOfLists[p[0:ind]].append(int(p[ind+1:]))
    
    for i in dictOfLists:
        dictOfLists[i] = sorted(dictOfLists[i])
    
    return dictOfLists, totcount_phase



def ranges(i):
    for a, b in itertools.groupby(enumerate(i), lambda pair: pair[1] - pair[0]):
        b = list(b)
        yield b[0][1], b[-1][1]



def rangesList(dictOfLists,countofPairs):

    for i in dictOfLists:
        dictOfLists[i] =  list(ranges(dictOfLists[i]))
        countofPairs += len(dictOfLists[i])
    return dictOfLists, countofPairs

        
        


countofPairs = 0
dictOfLists, totcount_phase = getDDict(tipomod)
list_of_phase, countofPairs = rangesList(dictOfLists,countofPairs)








def add_Title(document,listofdatt):
    paragrap = document.add_paragraph("REGLAMENTO DE COPROPIEDAD",style='bca')
    paragrap = document.add_paragraph("DEL "+str(pH),style='bca')
    paragrap = document.add_paragraph("CAPITULO I",style='bcal')
    paragrap = document.add_paragraph(str())
    paragrap = document.add_paragraph(str("DISPOSICIONES GENERALES"),style='bcal')
    if faSe != 1:
        paragrap = document.add_paragraph(str("LOS ARTICULOS 1, 2 Y 8 QUEDARAN DE LA SIGUIENTE FORMA:"),style='bcal')
    
    
def add_Art_1(document,listofdatt):
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    run = paragraph.add_run("ARTICULO 1: ")
    run.underline = True
    run.bold = True
    paragraph.add_run("Quedan sometidas al Régimen de Propiedad Horizontal con arreglo a las Disposiciones Legales previstas en la Ley doscientos ochenta y cuatro (284) del catorce (14) de febrero de dos mil veintidós (2022) (en adelante La Ley) y demás disposiciones pertinentes, la Finca inscrita al folio real  número ")
    run = paragraph.add_run(get_numbandword(str(int(NdeFinca))))
    run.bold = True
    paragraph.add_run(" con código de ubicación ")
    run = paragraph.add_run(get_numbandword(str(int(codUb))))
    run.bold = True
    paragraph.add_run(", ubicada en el ")
    
    run = paragraph.add_run(str(ubDeCod))
    run.bold = True
    paragraph.add_run(", inscrita la sección de la Propiedad, del Registro Público, Provincia de Panamá, así como las mejoras construidas sobre dicha finca; en adelante, denominada ")
    run = paragraph.add_run(str(pH)+".")
    run.bold = True

def add_Art_2(document,listofdatt):
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    paragraph.alignment = 3
    run = paragraph.add_run("ARTICULO 2: ")
    run.underline = True
    run.bold = True
    paragraph.add_run("EL ")
    run = paragraph.add_run(str(pH))
    run.bold = True
    paragraph.add_run(" está ubicado en ")
    run = paragraph.add_run(str(ubGeneral))
    run.bold = True
    paragraph.add_run(", y la conforma un lote de terreno con una superficie de ")
    run = paragraph.add_run(get_numbandword(str(aTotTerr)))
    run.bold = True
    paragraph.add_run(" y un valor de ")
    run = paragraph.add_run(get_numbandword(str(valTerr)))
    run.bold = True
    paragraph.add_run(".")
    
    
    
    document.add_paragraph("")
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3

    paragraph.add_run("Sobre ella se ha construido un conjunto de unidades inmobiliarias residenciales, el cual consiste en un desarrollo urbanístico residencial, por lo cual será de ")
    run = paragraph.add_run("USO RESIDENCIAL")
    run.bold = True
    paragraph.add_run(", el conjunto de unidades inmobiliarias residenciales, ha sido construido con estructura de hormigón, paredes de concreto armado repelladas en ambas caras, piso de baldosa, ventana con marco de aluminio y vidrio corredizo, techo de fibrocemento o láminas de acero galvanizado (según el modelo), el cual será denominado en su conjunto como ")
    run = paragraph.add_run(str(pH))
    run.bold = True
    run = paragraph.add_run(".")
    run.bold = True
    
    
    document.add_paragraph("")
    paragraph = document.add_paragraph("La construcción del proyecto ocupa una superficie de ")
    paragraph.alignment = 3
    run = paragraph.add_run(get_numbandword(str(aTotTerr)))
    run.bold = True
    paragraph.add_run(",  de los cuales, ")
    run = paragraph.add_run(get_numbandword(str(aTotConst)))
    run.bold = True
    paragraph.add_run(" corresponden al área de construcción de las unidades inmobiliarias, ")
    run = paragraph.add_run(get_numbandword(str(aTotCom)))
    run.bold = True
    paragraph.add_run(" corresponden a las áreas comunes entrada vehicular y peatonal, garita de entrada, calles y vía de acceso, aceras, área de parque y áreas verdes")

    if str(aResComer) != " mts2" and str(aResComer) != "0.00"and aResComer != "0 mts2":
        run = paragraph.add_run(" y "+get_numbandword(str(aResComer)))
        run.bold = True
        paragraph.add_run(" serán reserva del promotor, quedando en la finca propiedad ")
        run = paragraph.add_run(get_numbandword(str(int(NdeFinca))))
        run.bold = True
        paragraph.add_run(", con código de ubicación ")
        run = paragraph.add_run(get_numbandword(str(int(codUb))))
        run.bold = True
    
    paragraph.add_run(".")
    
    
    
def add_Desc_Gen(document,listofdatt,dicTOT):


    paragraph = document.add_paragraph("DESCRIPCION GENERAL DEL ",style='ba')
    paragraph.add_run(str(pH)+".")
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    run = paragraph.add_run(str(descrGeneral))
    document.add_paragraph("")
    
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3


    
    
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    run = paragraph.add_run("El ")
    run = paragraph.add_run(str(pH)+",")
    run.bold = True
    run = paragraph.add_run(" estará compuesto por ")
    run = paragraph.add_run(get_numbandword(str(int(uniTotal))))
    run = paragraph.add_run(" unidades inmobiliarias en total, las cuales serán identificadas de la ")
    run = paragraph.add_run(manzana+"-001 a la "+manzana+"-"+str(int(uniTotal)) + ".")
        
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    run = paragraph.add_run("El ")
    run = paragraph.add_run(str(pH)+",")
    run.bold = True
    run = paragraph.add_run(" estará compuesto por ")
    run = paragraph.add_run(get_numbandword(str(int(uniTotal))))
    run = paragraph.add_run(" unidades inmobiliarias, desglosadas así:")
    for i in dicTOT:
        run = paragraph.add_run(" "+get_numbandword(str(dicTOT[i])))
        run = paragraph.add_run(" unidades modelo ")
        run = paragraph.add_run(str(i)+",")

    run = paragraph.add_run(" las cuales se describen a continuación:")

    
    
def add_Desc_Rep(document,listofdatt,dicdeff):
    for i in dicTOT:
    
        paragraph = document.add_paragraph("UNIDAD INMOBILIARIA ",style='ba')
        x = i.split()
        if len(x)> 2:
            for p in range(len(x)-2):
                p +=2
                x[1] =x[1]+ " "+x[p]
        if len(x) == 1:
            run = paragraph.add_run(i+":")
        else:
            run = paragraph.add_run(x[0]+'"'+x[1]+'"'+":")
            
        paragraph = document.add_paragraph("")
        paragraph.alignment = 3
        run = paragraph.add_run("Descripción de la Unidad inmobiliaria: ")
        run.bold = True
        run.underline = True
        run = paragraph.add_run(dicdeff[i])
        document.add_paragraph("")
        
def add_Desc_Rep2(document,listofdatt,dicdeff, tipomod):
    for i in tipomod:
        if len(tipomod[i]) != 0:
            paragraph = document.add_paragraph("UNIDAD INMOBILIARIA ",style='ba')
            x = i.split()

            if len(x)> 2:
                for p in range(len(x)-2):
                    p +=2
                    x[1] =x[1]+ " "+x[p]



            if len(x) == 1:
                run = paragraph.add_run(i+":")
            else:
                run = paragraph.add_run(x[0]+'"'+x[1]+'"'+":")

            paragraph = document.add_paragraph("")
            paragraph.alignment = 3
            run = paragraph.add_run("Cantidad de Unidades Inmobiliarias y Nomenclatura: ")
            run.bold = True
            run.underline = True
            run = paragraph.add_run("Comprende de ")
            run = paragraph.add_run(get_numbandword( str(len(tipomod[i])  )))
            run = paragraph.add_run(" unidades inmobiliarias identificadas como:  ")
            ccc = 0
            for p in tipomod[i]:
                if ccc != len(tipomod[i])-1:
                    run = paragraph.add_run(p+", ")
                else:
                    run = paragraph.add_run(p+".")
                ccc += 1
            

            paragraph = document.add_paragraph("")
            paragraph.alignment = 3
            paragraph = document.add_paragraph("")
            paragraph.alignment = 3
            run = paragraph.add_run("Descripción de la Unidad inmobiliaria: ")
            run.bold = True
            run.underline = True
            run = paragraph.add_run(dicdeff[i])
            paragraph = document.add_paragraph("")
            paragraph.alignment = 3
    




def add_Desc_Etapa_Current(document,listofdatt,dict_letters,list_of_phase,totcount_phase,tipomod):
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    run = paragraph.add_run("La construcción de la ")
    run = paragraph.add_run(dictEnum[int(faSe)])
    run = paragraph.add_run(" etapa del proyecto ocupará una superficie de ")
    run = paragraph.add_run(get_numbandword(str(totFase)))
    run = paragraph.add_run(", de los cuales ")
    run = paragraph.add_run(get_numbandword(str(constFase)))
    run = paragraph.add_run(" corresponden al área de construcción de las unidades inmobiliarias y ")
    run = paragraph.add_run(get_numbandword(str(comFase)))
    run = paragraph.add_run("corresponden a áreas comunes que incluyen calles, gramas y aceras del proyecto.")
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    run = paragraph.add_run("La construcción de la ")
    run = paragraph.add_run(dictEnum[int(faSe)])
    run = paragraph.add_run(" etapa del proyecto estará compuesta por ")
    run = paragraph.add_run(get_numbandword(str(totcount_phase)))
    run = paragraph.add_run(" unidades inmobiliarias, las cuales serán identificadas como ")
    cccc = 0
    list_of_phase
    for p in list_of_phase:
        
        for i in list_of_phase[p]:
            i = list(i)
            run = paragraph.add_run(p+"-"+str(i[0])+" a la ")
            run = paragraph.add_run(p+"-"+str(i[1])+",")
            if cccc != countofPairs-1:
                run = paragraph.add_run(" y ")
            cccc += 1
    run = paragraph.add_run(" compuestas por ")
    
    dict_letters2 = {}
    for p in dict_letters:
        
        
        if len(dict_letters[p]) == 0:
            if p in tipomod and len(tipomod[p]) != 0:
                dict_letters2[p] = []
        else:
            
            
            for i in dict_letters[p]:
                if p+" "+i in tipomod and len(tipomod[p+" "+i]) != 0:
                    if p not in dict_letters2:
                        dict_letters2[p] = []
                    dict_letters2[p].append(i)
   
    for i in dict_letters2:
        run = paragraph.add_run("Unidades Inmobiliaria Tipo ")
        run = paragraph.add_run(i+" ")
        for l in dict_letters2[i]:
            run = paragraph.add_run('"'+l+'", ')
            
    run = paragraph.add_run(" los cuales se detallan a continuación:")
    
def add_Desc_Reserva(document,listofdatt,dicdeff):
    paragraph = document.add_paragraph("En las siguientes fases a construirse")
    paragraph.alignment = 3
    if str(aRes) != "0.00 mts2":
        run = paragraph.add_run(", sobre la reserva de ")
        run = paragraph.add_run(get_numbandword(str(aRes)) + ".")
        run.bold = True

    run = paragraph.add_run(" Se edificarán las mejoras cuando así lo determine ")
    run = paragraph.add_run("EL PROMOTOR,")
    run.bold = True
    run = paragraph.add_run(" las cuales, una vez estén construidas, se solicitará la aprobación al Ministerio de Vivienda y Ordenamiento Territorial, para que sean incorporadas y formen un todo con las mejoras de la Primera Etapa")
    run = paragraph.add_run(" del ")
    
    run = paragraph.add_run(str(pH)+".")
    run.bold = True
    
    
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    paragraph = document.add_paragraph("De esta manera, también es entendido que serán incorporados al presente Régimen de Propiedad Horizontal del ")
    paragraph.alignment = 3
    run = paragraph.add_run(str(pH))
    run.bold = True
    run = paragraph.add_run(" todos los lotes que correspondan a las futuras fases a construirse, para lo cual en su momento se aportarán los correspondientes planos que sustenten estos lotes, así como todas aquellas mejoras que el Promotor haya edificado o edifique en el futuro sobre los mismos, los cuales conformarán fincas registrales aparte como Bienes Privados.")
    
    
    
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    paragraph = document.add_paragraph("El Conjunto residencial será desarrollado por etapas, para lo cual, el promotor se reserva el derecho de edificar mejoras privadas en")
    paragraph.alignment = 3
    xrang = range(int(faSe))
    for i in xrang:
        run = paragraph.add_run(" la ")
       
        run = paragraph.add_run(str(dictEnum[int(i+1)])+" etapa,")

    run = paragraph.add_run(" y en las futuras Etapas que se construyan en el área de reserva del Proyecto, el cual cuenta con una superficie de ")
    run = paragraph.add_run(get_numbandword(str(aRes)) +', ')
    run.bold = True
    run = paragraph.add_run("las cuales serán incorporadas al Régimen de Propiedad Horizontal del ")
    run = paragraph.add_run(str(pH)+",")
    run.bold = True
    run = paragraph.add_run(" sin que para ello requiera del consentimiento de las autoridades del P.H., los propietarios, la Asamblea de Propietarios, la Junta Directiva y/o el Administrador del ")
    run = paragraph.add_run(str(pH)+". ")
    run.bold = True
    if str(aResComer) != " mts2" and str(aResComer) != "0.00"and aResComer != "0 mts2"and aResComer !=  '0.00 mts2':
        run = paragraph.add_run("los "+get_numbandword(str(aResComer))+" será resto libre, quedando en la finca propiedad "+str(int(NdeFinca))+", con código de ubicación "+str(int(codUb))+".")



    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    paragraph = document.add_paragraph("")
    paragraph.alignment = 3





    run = paragraph.add_run("EL PROMOTOR ")
    run.bold = True
    run = paragraph.add_run("se reserva el derecho de colocar cámaras de vigilancia en las áreas comunes del ")
    run = paragraph.add_run(str(pH)+",")
    run.bold = True
    run = paragraph.add_run(" las cuales serán manejadas en la garita de seguridad del ")
    run = paragraph.add_run(str(pH))
    run.bold = True
    run = paragraph.add_run(" y las imágenes captadas por la misma solo cumplirán funciones de vigilancia y serán facilitadas únicamente a las autoridades competentes al momento de ser requeridas por la misma.")



def write_artTemp(document,listofdatt):
    


    paragraph = document.add_paragraph("")
    paragraph.alignment = 3
    run = paragraph.add_run("ARTÍCULO 3: El")
    run.bold = True
    run = paragraph.add_run(artnum+": ")
    
    
  





def etapa_modul(document,listofdatt,dict_letters,list_of_phase,totcount_phase,dicdeff, tipomod):
    global faSe
    iiin = faSe

    for i in range(int(faSe)):
        i +=1
     
            
        dict_letters = {}
        listofdatt , tipomod , dicdeff, dicTOT, manzana= getVals(i)

        faSe = i
        

        global totFase, constFase, comFase, aRes
        totFase = str(MakeSQRmtrr(listofdatt[faSe][7]))+" mts2"
        constFase = str(MakeSQRmtrr(listofdatt[faSe][8]))+" mts2"
        comFase = str(MakeSQRmtrr(listofdatt[faSe][9]))+" mts2"
        aRes = str(MakeSQRmtrr(listofdatt[faSe][10]))+" mts2"



        for p in dicdeff:
            x = p.split()

            if len(x)> 2:
                for i in range(len(x)-2):
                    i +=2
                    x[1] =x[1]+ " "+x[i]
            if x[0] not in dict_letters:
                dict_letters[x[0]] = []
            
            if len(x) != 1:
                dict_letters[x[0]].append(x[1])
            

        countofPairs = 0
        

        dictOfLists, totcount_phase = getDDict(tipomod)

                
                



        list_of_phase, countofPairs = rangesList(dictOfLists,countofPairs)
        
        
        
        
        
        
        add_Desc_Etapa_Current(document,listofdatt,dict_letters,list_of_phase,totcount_phase,tipomod)
        paragrap = document.add_paragraph(str())
        add_Desc_Rep2(document,listofdatt,dicdeff, tipomod)
        paragrap = document.add_paragraph(str())
        
        
        if faSe == iiin:
            add_Desc_Reserva(document,listofdatt,dicdeff)
            paragrap = document.add_paragraph(str())



def CreateReg(etapa):    
    global faSe
    faSe = etapa
    document = docx.Document()
    document.styles['Normal'].font.name = 'Arial'

    style = document.styles.add_style('bca', WD_STYLE_TYPE.PARAGRAPH)
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = style.font
    font.bold = True
    font.name = 'Arial'

    style = document.styles.add_style('bcal', WD_STYLE_TYPE.PARAGRAPH)
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = style.font
    font.underline = True
    font.bold = True
    font.name = 'Arial'

    style = document.styles.add_style('ba', WD_STYLE_TYPE.PARAGRAPH)
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    font = style.font
    font.bold = True
    font.name = 'Arial'

    style = document.styles.add_style('bal', WD_STYLE_TYPE.PARAGRAPH)
    #style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    font = style.font
    font.underline = True
    font.bold = True
    font.name = 'Arial'
    add_Title(document,listofdatt)
    paragrap = document.add_paragraph(str())
    add_Art_1(document,listofdatt)
    paragrap = document.add_paragraph(str())
    add_Art_2(document,listofdatt)
    paragrap = document.add_paragraph(str())
    add_Desc_Gen(document,listofdatt,dicTOT)
    paragrap = document.add_paragraph(str())
    add_Desc_Rep(document,listofdatt,dicdeff)
    paragrap = document.add_paragraph(str())
    etapa_modul(document,listofdatt,dict_letters,list_of_phase,totcount_phase,dicdeff, tipomod)
    document.save("Output_Docx/Parts/part1.docx")   



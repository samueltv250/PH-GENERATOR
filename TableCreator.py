from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from extract_Alcance_PH import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from numero_letras import *


def addHeader(table,llist,numb):
    dictH = {1:"",2:" SEGUNDA II ETAPA",3:" TERCERA III ETAPA",4:" CUARTA IV ETAPA",5:" QUINTA V ETAPA"}
    row = table.add_row().cells
    a = row[0]
    b = row[4]
    A = a.merge(b)
    p = A.add_paragraph("")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(llist[0][0]+dictH[numb])
    run.bold = True
    font = run.font
    font.name = 'Arial'
    
    
    row = table.add_row().cells
    parag = row[0].add_paragraph("")
    run = parag.add_run("UNIDAD")
    run.bold = True
    parag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = run.font
    font.name = 'Arial'

   
    parag = row[1].add_paragraph("")
    run = parag.add_run("VALOR DE TERRENO")
    run.bold = True
    parag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = run.font
    font.name = 'Arial'


  
    parag = row[2].add_paragraph("")
    run = parag.add_run("VALOR DE MEJORAS")
    run.bold = True
    parag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = run.font
    font.name = 'Arial'

 
    parag = row[3].add_paragraph("")
    run = parag.add_run("VALOR TOTAL")
    run.bold = True
    parag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = run.font
    font.name = 'Arial'


 
    parag = row[4].add_paragraph("")
    run = parag.add_run('%')
    run.bold = True
    parag.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = run.font
    font.name = 'Arial'
def merge_two_dicts(x, y):
    z = x.copy()   # start with keys and values of x
    z.update(y)    # modifies z with keys and values of y
    return z

def requestCreator():
    data = TablaFinal()
    mainDICT = extractDictEt_mod()
    dictAreasCerr,dictAreasAbiert,dictAreasPav = extractTableDat()
    RequestDict = {}
    z = {}
    for i in mainDICT:
        z = merge_two_dicts(z, mainDICT[i])
    tabldict = {}
    for i in data:
        tabl = data[i]
        for p in tabl:
            tabldict[p[0]] = p[1:]
        


        

    for  i in z:
        moddel = z[i]
        if moddel not in RequestDict:
            RequestDict[moddel] = []
        RequestDict[moddel].append([i,dictAreasCerr[moddel],moddel,tabldict[i][0],tabldict[i][1],tabldict[i][2],tabldict[i][3]])
    return RequestDict
import pandas as pd
import xlsxwriter
def tableToModelT(RequestDict):
   
    writer = pd.ExcelWriter('Output_Request/Request.xlsx', engine='xlsxwriter')

    

    for model in RequestDict:
        llist = [["No. Lote","M2","TIPO","VALOR UNITARIO B/.","VALOR DE TERRENO","VALOR TOTAL","PORCENTAJE"]]
        for col in RequestDict[model]:
            llist.append(col)
        df = pd.DataFrame(llist)
        df.to_excel(writer, sheet_name=str(model), index=False)
        # print(llist)
    writer.save()
        


        
def fullReq():
    requestDict = requestCreator()
    tableToModelT(requestDict)







# fullReq()

def addEtap(table,numb):
    data = TablaFinal()[numb]
    totEtap = [0,0,0,0]
    dictnum = {1:"I",2:"II",3:"III",4:"IV",5:"V",6:"VI"}
    for unidad,valmej, valterr,valTot,porcent in data:
        valterr = valterr
        valmej= valmej
        valTot= valTot
        porcent= porcent
   
        row = table.add_row().cells
        row[0].text = str(unidad)
        row[1].text = MakeSQRmtrr(valterr)
        totEtap[0] += valterr
        row[2].text = MakeSQRmtrr(valmej)
        totEtap[1] += valmej
        row[3].text = MakeSQRmtrr(valTot)
        totEtap[2] += valTot
        row[4].text = MakeSQRmtrr(porcent)
        totEtap[3] += porcent
    row = table.add_row().cells
    row[0].text = str("TOTAL "+dictnum[numb]+" ETAPA")
    row[1].text = MakeSQRmtrr(totEtap[0])
    row[2].text = MakeSQRmtrr(totEtap[1])
    row[3].text = MakeSQRmtrr(totEtap[2])
    row[4].text = MakeSQRmtrr(totEtap[3])
    return totEtap


def createTabl(numb):
    document = Document()
    totmnrest,restolibre,AREAconst =Variabl()
    llist = listOfDat()
    totPrice = llist[0][2]
    document = Document()
    table = document.add_table(rows=0, cols=5)
    table.style = 'Table Grid'
    totetP =[0,0,0,0]
    
    for o in range(numb):
        addHeader(table,llist,o+1)
        temp = addEtap(table,o+1)

        for i in range(len(temp)):
            totetP[i] += temp[i]
    
    reserv = totPrice-totetP[0]-restolibre
    if reserv != 0 and reserv != str(0)and reserv != str("0,00")and reserv != str("0.00"):
        row = table.add_row().cells
        parag = row[0].add_paragraph("")
        run = parag.add_run("RESERVA")
        run.bold = True
        parag = row[1].add_paragraph("")
        run = parag.add_run(MakeSQRmtrr(reserv))
        run.bold = True
        parag = row[2].add_paragraph("")
        run = parag.add_run("-------------------")
        run.bold = True

        parag = row[3].add_paragraph("")
        run = parag.add_run(MakeSQRmtrr(reserv))
        run.bold = True

        parag = row[4].add_paragraph("")
        run = parag.add_run(MakeSQRmtrr(100-totetP[3]))
        run.bold = True
   
        totetP[0] += reserv
        totetP[2] += reserv
        totetP[3] = 100
    if restolibre != 0 and restolibre != str(0)and restolibre != str("0,00")and reserv != str("0.00"):
        row = table.add_row().cells
        parag = row[0].add_paragraph("")
        run = parag.add_run("RESTOLIBRE")
        run.bold = True
        parag = row[1].add_paragraph("")
        run = parag.add_run(MakeSQRmtrr(restolibre))
        run.bold = True
        parag = row[2].add_paragraph("")
        run = parag.add_run("-------------------")
        run.bold = True
        parag = row[3].add_paragraph("")
        run = parag.add_run(MakeSQRmtrr(restolibre))
        run.bold = True
        parag = row[4].add_paragraph("")
        run = parag.add_run("-------------------")
        run.bold = True
        totetP[0] += restolibre
        totetP[2] += restolibre
    row = table.add_row().cells
    parag = row[0].add_paragraph("")
    run = parag.add_run("TOTAL")
    run.bold = True
    parag = row[1].add_paragraph("")
    run = parag.add_run(MakeSQRmtrr(totetP[0]))
    run.bold = True
    parag = row[2].add_paragraph("")
    run = parag.add_run(MakeSQRmtrr(totetP[1]))
    run.bold = True
    parag = row[3].add_paragraph("")
    run = parag.add_run(MakeSQRmtrr(totetP[2]))
    run.bold = True

    parag = row[4].add_paragraph("")
    run = parag.add_run(MakeSQRmtrr(totetP[3]))
    run.bold = True
 

    
    
    document.save("Output_Docx/Parts/table.docx")



